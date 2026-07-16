import hashlib
import json
import unicodedata
from abc import ABC, abstractmethod
from collections.abc import Iterable, Mapping
from datetime import date, datetime, time
from decimal import Decimal
from io import BytesIO
from pathlib import Path
from typing import BinaryIO
from urllib.parse import quote

import pandas as pd
from docx import Document
from pypdf import PdfReader

from app.etl.dataframe import records_to_dataframe
from app.etl.models import CleanStatus, DocumentRecord, JsonValue, Metadata


class FileLoadError(ValueError):
    """Raised when a supported source cannot be extracted safely."""


def normalized_content_hash(content: str) -> str:
    """Hash content after stable Unicode and newline normalization."""
    normalized = unicodedata.normalize("NFC", content).replace("\r\n", "\n").replace("\r", "\n")
    return hashlib.sha256(normalized.encode("utf-8")).hexdigest()


def _source_metadata(
    metadata: Mapping[str, JsonValue] | None,
    *,
    source_type: str,
    source_key: str,
    **specific: JsonValue,
) -> Metadata:
    result = dict(metadata or {})
    result.update(specific)
    result["source_type"] = source_type
    result["location"] = source_key
    return result


def _record(content: str, *, source_key: str, metadata: Metadata) -> DocumentRecord:
    return DocumentRecord(
        content=content,
        metadata=metadata,
        source_key=source_key,
        content_hash=normalized_content_hash(content),
        clean_status=CleanStatus.CLEANED,
    )


def _read_bytes(source: BinaryIO) -> bytes:
    try:
        source.seek(0)
        return source.read()
    except (OSError, ValueError) as exc:
        raise FileLoadError("source stream could not be read") from exc


def _decode_text(source: BinaryIO, *, source_type: str) -> str:
    try:
        return _read_bytes(source).decode("utf-8-sig")
    except UnicodeDecodeError as exc:
        raise FileLoadError(f"{source_type} source must be UTF-8 encoded") from exc


def _json_value(value: object) -> JsonValue:
    if value is None or value is pd.NA:
        return None
    if isinstance(value, datetime | date | time):
        return value.isoformat()
    if isinstance(value, Decimal):
        return int(value) if value == value.to_integral_value() else float(value)
    if isinstance(value, str | int | float | bool):
        if isinstance(value, float) and pd.isna(value):
            return None
        return value
    if pd.isna(value):
        return None
    return str(value)


def _structured_content(row_data: Mapping[str, JsonValue]) -> str:
    return json.dumps(row_data, ensure_ascii=False, separators=(",", ":"), sort_keys=True)


class FileLoader(ABC):
    """Base adapter exposing records and their canonical DataFrame representation."""

    source_type: str

    @abstractmethod
    def load(
        self,
        source: BinaryIO,
        *,
        source_key: str,
        metadata: Mapping[str, JsonValue] | None = None,
    ) -> Iterable[DocumentRecord]:
        """Extract canonical records without applying business cleaning."""

    def load_dataframe(
        self,
        source: BinaryIO,
        *,
        source_key: str,
        metadata: Mapping[str, JsonValue] | None = None,
    ) -> pd.DataFrame:
        return records_to_dataframe(self.load(source, source_key=source_key, metadata=metadata))


class TextLoader(FileLoader):
    source_type = "TXT"

    def load(
        self,
        source: BinaryIO,
        *,
        source_key: str,
        metadata: Mapping[str, JsonValue] | None = None,
    ) -> Iterable[DocumentRecord]:
        content = _decode_text(source, source_type=self.source_type)
        yield _record(
            content,
            source_key=source_key,
            metadata=_source_metadata(
                metadata, source_type=self.source_type, source_key=source_key
            ),
        )


class MarkdownLoader(TextLoader):
    source_type = "MD"


class PdfLoader(FileLoader):
    source_type = "PDF"

    def load(
        self,
        source: BinaryIO,
        *,
        source_key: str,
        metadata: Mapping[str, JsonValue] | None = None,
    ) -> Iterable[DocumentRecord]:
        try:
            reader = PdfReader(BytesIO(_read_bytes(source)))
            if reader.is_encrypted and reader.decrypt("") == 0:
                raise FileLoadError("PDF source is encrypted")
            pages = [(page.extract_text() or "") for page in reader.pages]
        except FileLoadError:
            raise
        except Exception as exc:
            raise FileLoadError("PDF source could not be parsed") from exc

        for page_number, content in enumerate(pages, start=1):
            record_key = f"{source_key}#page={page_number}"
            yield _record(
                content,
                source_key=record_key,
                metadata=_source_metadata(
                    metadata,
                    source_type=self.source_type,
                    source_key=source_key,
                    page=page_number,
                ),
            )


class DocxLoader(FileLoader):
    source_type = "DOCX"

    def load(
        self,
        source: BinaryIO,
        *,
        source_key: str,
        metadata: Mapping[str, JsonValue] | None = None,
    ) -> Iterable[DocumentRecord]:
        try:
            document = Document(BytesIO(_read_bytes(source)))
            paragraphs = [paragraph.text for paragraph in document.paragraphs]
            table_rows = [
                "\t".join(cell.text for cell in row.cells)
                for table in document.tables
                for row in table.rows
            ]
        except Exception as exc:
            raise FileLoadError("DOCX source could not be parsed") from exc

        content = "\n".join([*paragraphs, *table_rows])
        yield _record(
            content,
            source_key=source_key,
            metadata=_source_metadata(
                metadata, source_type=self.source_type, source_key=source_key
            ),
        )


class CsvLoader(FileLoader):
    source_type = "CSV"

    def load(
        self,
        source: BinaryIO,
        *,
        source_key: str,
        metadata: Mapping[str, JsonValue] | None = None,
    ) -> Iterable[DocumentRecord]:
        try:
            frame = pd.read_csv(BytesIO(_read_bytes(source)), encoding="utf-8-sig", dtype=object)
        except (UnicodeDecodeError, pd.errors.ParserError) as exc:
            raise FileLoadError("CSV source could not be parsed as UTF-8") from exc
        except pd.errors.EmptyDataError:
            return

        yield from _tabular_records(
            frame,
            source_key=source_key,
            source_type=self.source_type,
            metadata=metadata,
        )


class XlsxLoader(FileLoader):
    source_type = "XLSX"

    def load(
        self,
        source: BinaryIO,
        *,
        source_key: str,
        metadata: Mapping[str, JsonValue] | None = None,
    ) -> Iterable[DocumentRecord]:
        try:
            sheets = pd.read_excel(
                BytesIO(_read_bytes(source)), sheet_name=None, dtype=object, engine="openpyxl"
            )
        except Exception as exc:
            raise FileLoadError("XLSX source could not be parsed") from exc

        for sheet_name, frame in sheets.items():
            yield from _tabular_records(
                frame,
                source_key=source_key,
                source_type=self.source_type,
                metadata=metadata,
                sheet=str(sheet_name),
            )


def _tabular_records(
    frame: pd.DataFrame,
    *,
    source_key: str,
    source_type: str,
    metadata: Mapping[str, JsonValue] | None,
    sheet: str | None = None,
) -> Iterable[DocumentRecord]:
    columns = [str(column) for column in frame.columns]
    for zero_based_row, values in enumerate(frame.itertuples(index=False, name=None)):
        row_number = zero_based_row + 2
        row_data = {
            column: _json_value(value) for column, value in zip(columns, values, strict=True)
        }
        suffix = f"#row={row_number}"
        specific: dict[str, JsonValue] = {"row": row_number, "row_data": row_data}
        if sheet is not None:
            suffix = f"#sheet={quote(sheet, safe='')}&row={row_number}"
            specific["sheet"] = sheet
        yield _record(
            _structured_content(row_data),
            source_key=f"{source_key}{suffix}",
            metadata=_source_metadata(
                metadata,
                source_type=source_type,
                source_key=source_key,
                **specific,
            ),
        )


_LOADERS: dict[str, type[FileLoader]] = {
    ".csv": CsvLoader,
    ".docx": DocxLoader,
    ".md": MarkdownLoader,
    ".pdf": PdfLoader,
    ".txt": TextLoader,
    ".xlsx": XlsxLoader,
}


def loader_for(source_key: str) -> FileLoader:
    """Select a loader from a source key's case-insensitive file extension."""
    suffix = Path(source_key).suffix.lower()
    try:
        return _LOADERS[suffix]()
    except KeyError as exc:
        supported = ", ".join(sorted(_LOADERS))
        raise FileLoadError(
            f"unsupported file extension {suffix or '<none>'}; expected {supported}"
        ) from exc
