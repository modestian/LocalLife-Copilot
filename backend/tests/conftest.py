from io import BytesIO

import pandas as pd
import pytest
from docx import Document
from pypdf import PdfWriter


@pytest.fixture(scope="session")
def ingestion_format_payloads() -> dict[str, bytes]:
    """Small, deterministic source files covering every supported ingestion format."""
    docx_stream = BytesIO()
    document = Document()
    document.add_paragraph("安静舒适")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "人均"
    table.rows[0].cells[1].text = "45 元"
    document.save(docx_stream)

    pdf_stream = BytesIO()
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    writer.write(pdf_stream)

    xlsx_stream = BytesIO()
    with pd.ExcelWriter(xlsx_stream, engine="openpyxl") as excel_writer:
        pd.DataFrame([{"商家": "青禾", "评分": 4.8}]).to_excel(
            excel_writer, sheet_name="门店 信息", index=False
        )
        pd.DataFrame([{"菜品": "拿铁"}]).to_excel(excel_writer, sheet_name="菜单", index=False)

    return {
        "sample.pdf": pdf_stream.getvalue(),
        "sample.docx": docx_stream.getvalue(),
        "sample.md": "# 推荐\n\n环境安静。".encode(),
        "sample.txt": "环境安静\n适合聊天。".encode(),
        "sample.csv": "商家,评分,备注\n青禾,4.8,\n木棉,5,安静\n".encode("utf-8-sig"),
        "sample.xlsx": xlsx_stream.getvalue(),
    }


@pytest.fixture(scope="session")
def invalid_ingestion_payloads() -> dict[str, bytes]:
    """Malformed sources whose suffix selects a loader but whose bytes are invalid."""
    return {
        "invalid.pdf": b"not a PDF",
        "invalid.docx": b"not a DOCX archive",
        "invalid.txt": b"\xff",
        "invalid.md": b"\xff",
        "invalid.csv": b"\xff",
        "invalid.xlsx": b"not an XLSX archive",
    }
