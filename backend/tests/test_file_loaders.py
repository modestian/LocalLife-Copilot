from io import BytesIO

import pandas as pd
import pytest
from docx import Document
from pypdf import PdfWriter

from app.etl.dataframe import CANONICAL_COLUMNS
from app.etl.loaders import (
    CsvLoader,
    DocxLoader,
    FileLoadError,
    MarkdownLoader,
    PdfLoader,
    TextLoader,
    XlsxLoader,
    loader_for,
    normalized_content_hash,
)
from app.etl.models import CleanStatus


def docx_source() -> BytesIO:
    stream = BytesIO()
    document = Document()
    document.add_paragraph("安静舒适")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "人均"
    table.rows[0].cells[1].text = "45 元"
    document.save(stream)
    stream.seek(0)
    return stream


def pdf_source() -> BytesIO:
    stream = BytesIO()
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    writer.write(stream)
    stream.seek(0)
    return stream


def xlsx_source() -> BytesIO:
    stream = BytesIO()
    with pd.ExcelWriter(stream, engine="openpyxl") as writer:
        pd.DataFrame([{"商家": "青禾", "评分": 4.8}]).to_excel(
            writer, sheet_name="门店 信息", index=False
        )
        pd.DataFrame([{"菜品": "拿铁"}]).to_excel(writer, sheet_name="菜单", index=False)
    stream.seek(0)
    return stream


@pytest.mark.parametrize(
    ("loader", "suffix", "content", "source_type"),
    [
        (TextLoader(), "txt", "环境安静\r\n适合聊天。", "TXT"),
        (MarkdownLoader(), "md", "# 推荐\n\n环境安静。", "MD"),
    ],
)
def test_text_loaders_extract_one_canonical_record(
    loader: TextLoader, suffix: str, content: str, source_type: str
) -> None:
    records = list(
        loader.load(
            BytesIO(content.encode("utf-8-sig")),
            source_key=f"reviews/example.{suffix}",
            metadata={"merchant_id": 7},
        )
    )

    assert len(records) == 1
    assert records[0].content == content
    assert records[0].source_key == f"reviews/example.{suffix}"
    assert records[0].metadata == {
        "merchant_id": 7,
        "source_type": source_type,
        "location": f"reviews/example.{suffix}",
    }
    assert records[0].clean_status is CleanStatus.CLEANED
    assert records[0].content_hash == normalized_content_hash(content)


def test_docx_loader_extracts_paragraphs_and_tables() -> None:
    [record] = list(DocxLoader().load(docx_source(), source_key="guides/store.docx"))

    assert record.content == "安静舒适\n人均\t45 元"
    assert record.metadata["source_type"] == "DOCX"


def test_pdf_loader_preserves_page_boundaries_in_records() -> None:
    [record] = list(PdfLoader().load(pdf_source(), source_key="guides/store.pdf"))

    assert record.content == ""
    assert record.source_key == "guides/store.pdf#page=1"
    assert record.metadata["page"] == 1
    assert record.metadata["location"] == "guides/store.pdf"


def test_csv_loader_emits_json_compatible_rows_with_stable_keys() -> None:
    source = BytesIO("商家,评分,备注\n青禾,4.8,\n木棉,5,安静\n".encode("utf-8-sig"))

    records = list(CsvLoader().load(source, source_key="imports/shops.csv"))

    assert [record.source_key for record in records] == [
        "imports/shops.csv#row=2",
        "imports/shops.csv#row=3",
    ]
    assert records[0].content == '{"商家":"青禾","备注":null,"评分":"4.8"}'
    assert records[0].metadata["row_data"] == {"商家": "青禾", "评分": "4.8", "备注": None}
    assert records[1].metadata["row"] == 3


def test_xlsx_loader_emits_rows_for_every_sheet() -> None:
    records = list(XlsxLoader().load(xlsx_source(), source_key="imports/shops.xlsx"))

    assert [record.source_key for record in records] == [
        "imports/shops.xlsx#sheet=%E9%97%A8%E5%BA%97%20%E4%BF%A1%E6%81%AF&row=2",
        "imports/shops.xlsx#sheet=%E8%8F%9C%E5%8D%95&row=2",
    ]
    assert [record.metadata["sheet"] for record in records] == ["门店 信息", "菜单"]
    assert records[0].metadata["row_data"] == {"商家": "青禾", "评分": 4.8}


def test_load_dataframe_uses_documented_schema_and_types() -> None:
    frame = TextLoader().load_dataframe(BytesIO(b"example"), source_key="example.txt")

    assert tuple(frame.columns) == CANONICAL_COLUMNS
    assert frame.loc[0, "content"] == "example"
    assert frame.loc[0, "clean_status"] == "CLEANED"
    assert isinstance(frame.loc[0, "metadata"], dict)
    assert all(
        str(frame[column].dtype) == "string" for column in CANONICAL_COLUMNS if column != "metadata"
    )


@pytest.mark.parametrize(
    ("source_key", "expected_type"),
    [
        ("a.PDF", PdfLoader),
        ("a.docx", DocxLoader),
        ("a.MD", MarkdownLoader),
        ("a.txt", TextLoader),
        ("a.CSV", CsvLoader),
        ("a.xlsx", XlsxLoader),
    ],
)
def test_loader_for_selects_all_supported_extensions(
    source_key: str, expected_type: type[object]
) -> None:
    assert isinstance(loader_for(source_key), expected_type)


def test_loader_for_rejects_unsupported_extension() -> None:
    with pytest.raises(FileLoadError, match="unsupported file extension"):
        loader_for("imports/data.json")


def test_text_loader_rejects_non_utf8_input() -> None:
    with pytest.raises(FileLoadError, match="UTF-8"):
        list(TextLoader().load(BytesIO(b"\xff"), source_key="invalid.txt"))


def test_normalized_hash_is_stable_across_unicode_and_newline_forms() -> None:
    assert normalized_content_hash("cafe\u0301\r\n") == normalized_content_hash("caf\u00e9\n")
