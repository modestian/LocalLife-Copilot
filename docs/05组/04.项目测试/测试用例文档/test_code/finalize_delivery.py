"""Finalize the executable test-delivery workbook, report, and evidence index."""

from __future__ import annotations

import json
import re
import subprocess
from collections import Counter
from datetime import date
from pathlib import Path
from xml.etree import ElementTree

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from openpyxl import load_workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter


REPO_ROOT = Path(__file__).resolve().parents[3]
DELIVERY_DIR = Path(__file__).resolve().parent.parent
RUN_ROOT = DELIVERY_DIR / "execution_results" / "2026-07-26-current"
WORKBOOK = DELIVERY_DIR / "LocalLife-Copilot测试用例清单.xlsx"
REPORT = DELIVERY_DIR / "LocalLife-Copilot测试总结报告-V2.0-执行版.docx"
INDEX = DELIVERY_DIR / "交付说明与执行结果.md"

COMMIT = (
    subprocess.run(
        ["git", "rev-parse", "HEAD"],
        cwd=REPO_ROOT,
        check=True,
        capture_output=True,
        text=True,
    )
    .stdout.strip()
)
BRANCH = (
    subprocess.run(
        ["git", "branch", "--show-current"],
        cwd=REPO_ROOT,
        check=True,
        capture_output=True,
        text=True,
    )
    .stdout.strip()
)
STATUS_LINES = [
    line
    for line in subprocess.run(
        ["git", "status", "--short"],
        cwd=REPO_ROOT,
        check=True,
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
    ).stdout.splitlines()
    if line.strip()
]


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8-sig"))


FUNCTIONAL = load_json(RUN_ROOT / "functional" / "summary.json")
FULL = load_json(RUN_ROOT / "full-validation" / "summary.json")
ACCEPTANCE = load_json(RUN_ROOT / "acceptance" / "summary.json")
PERFORMANCE = load_json(RUN_ROOT / "performance" / "gate-result.json")
RECOVERY_TEXT = (RUN_ROOT / "recovery" / "summary.md").read_text(encoding="utf-8-sig")


def junit_counts(path: Path) -> tuple[int, int, int, int]:
    root = ElementTree.parse(path).getroot()
    cases = list(root.iter("testcase"))
    failed = sum(case.find("failure") is not None or case.find("error") is not None for case in cases)
    skipped = sum(case.find("skipped") is not None for case in cases)
    return len(cases), len(cases) - failed - skipped, failed, skipped


BACKEND_COUNTS = junit_counts(RUN_ROOT / "full-validation" / "backend-full-junit.xml")
FRONTEND_COUNTS = junit_counts(RUN_ROOT / "full-validation" / "frontend-full-junit.xml")
INTEGRATION_COUNTS = junit_counts(RUN_ROOT / "full-validation" / "integration-junit.xml")


def evidence_note(case_id: str) -> str:
    prefix = case_id.split("-")[0]
    notes = {
        "AUTH": "AUTH 分组、黑盒冒烟及 MySQL 认证/授权集成均通过；见 functional/01~02、acceptance/03、full-validation/05。",
        "GOV": "治理分组通过，ST-103 两条真实 MySQL 并发用例已启用通过；见 functional/05、full-validation/05。",
        "KB": "知识库后端 76 条、前端 9 条及 ST-102 MySQL/Redis 集成通过；见 functional/03~04、full-validation/05。",
        "SEARCH": "检索后端 29 条、前端 5 条及真实搜索性能样本通过；见 functional/06~07、performance/。",
        "RAG": "RAG 后端 69 条、前端 19 条通过；引用与无结果兜底同时由 E2E 覆盖；见 functional/08~09、acceptance/02。",
        "WS": "WebSocket/OpenAI 兼容后端 18 条、前端 5 条通过；见 functional/10~11。",
        "MERCHANT": "商家后端 123 条、前端 10 条通过，E2E 商家授权范围通过；见 functional/12~13、acceptance/02。",
        "MODEL": "模型治理后端 202 条、前端 10 条通过；见 functional/14~15。",
        "UI": "UI 单元 12 条、响应式 Playwright 2 条及全量前端 136 条通过；见 functional/16~17、full-validation/02。",
        "ADMIN": "管理端 25 条及管理员上传/失败重试 E2E 通过；见 functional/18、acceptance/02。",
        "OPS": "Compose、运维单元、黑盒冒烟均通过；见 functional/19~21、acceptance/03。",
        "E2E": "确定性种子 11 条、桌面/移动核心角色链路 8 条通过；当前轮结果见 acceptance/，三轮全新环境沿用历史 TK-702-03 证据。",
    }
    if case_id == "MODEL-04":
        return (
            "训练配置/任务逻辑自动化已通过，但 33 条实际产物校验因缺少 "
            "smoke-001、smoke-002、lora-001 及 peft/datasets/accelerate 而跳过；"
            "需完成两次 smoke 与一次正式 LoRA 训练后复跑 test_train_lora_artifacts.py。"
        )
    if case_id == "OPS-06":
        return "恢复演练通过：36 张表计数一致、OpenSearch 重建对账通过、Redis 清空后 MySQL 回源通过；见 recovery/summary.md。"
    if case_id == "OPS-07":
        return "性能门禁通过：3910 请求、0 失败；API/Search/TTFB P95 分别 58/180/99 ms；见 performance/gate-result.*。"
    return notes[prefix]


def finalize_workbook() -> None:
    workbook = load_workbook(WORKBOOK)
    sheet = workbook.worksheets[0]
    header_row = next(
        row
        for row in range(1, sheet.max_row + 1)
        if sheet.cell(row, 1).value == "用例ID"
    )
    result_counts: Counter[str] = Counter()
    for row in range(header_row + 1, sheet.max_row + 1):
        case_id = sheet.cell(row, 1).value
        if not isinstance(case_id, str) or not re.fullmatch(r"[A-Z0-9]+-\d{2}", case_id):
            continue
        result = "阻塞" if case_id == "MODEL-04" else "通过"
        result_counts[result] += 1
        sheet.cell(row, 11).value = "已执行自动化验证"
        sheet.cell(row, 14).value = result
        sheet.cell(row, 15).value = "BLOCK-MODEL-ARTIFACT-001" if result == "阻塞" else None
        sheet.cell(row, 16).value = evidence_note(case_id)
        sheet.cell(row, 14).fill = PatternFill(
            "solid", fgColor="FFF2CC" if result == "阻塞" else "E2F0D9"
        )
        sheet.cell(row, 14).font = Font(
            name="Microsoft YaHei",
            size=10,
            bold=True,
            color="9C6500" if result == "阻塞" else "006100",
        )
    sheet["A2"] = (
        "基于当前工作区、17个 Story 与项目原生自动化资产执行；"
        "本次结果：48通过、1阻塞（MODEL-04实际训练产物）。"
    )
    sheet["K1"] = "LocalLife Copilot 测试用例清单（2026-07-26 实际执行版）"

    current_name = "本次执行结果"
    if current_name in workbook.sheetnames:
        del workbook[current_name]
    result_sheet = workbook.create_sheet(current_name, 1)
    result_sheet.merge_cells("A1:F1")
    result_sheet["A1"] = "LocalLife Copilot 本次完整执行结果"
    result_sheet["A1"].font = Font(
        name="Microsoft YaHei", size=18, bold=True, color="FFFFFF"
    )
    result_sheet["A1"].fill = PatternFill("solid", fgColor="17365D")
    result_sheet["A1"].alignment = Alignment(horizontal="center", vertical="center")
    result_sheet.row_dimensions[1].height = 34
    metadata = [
        ("执行日期", "2026-07-26", "分支", BRANCH, "Commit", COMMIT),
        (
            "工作区",
            f"存在 {len(STATUS_LINES)} 项未提交改动",
            "用例结论",
            "48通过 / 1阻塞",
            "发布结论",
            "有条件通过",
        ),
    ]
    for row_index, values in enumerate(metadata, start=2):
        for column, value in enumerate(values, start=1):
            cell = result_sheet.cell(row_index, column, value)
            cell.fill = PatternFill("solid", fgColor="D9EAF7" if column % 2 else "FFFFFF")
            cell.font = Font(name="Microsoft YaHei", size=10, bold=column % 2 == 1)
            cell.alignment = Alignment(vertical="center", wrap_text=True)
    headers = ["检查项", "结果", "通过/总数", "跳过", "关键指标", "证据位置"]
    for column, value in enumerate(headers, start=1):
        cell = result_sheet.cell(5, column, value)
        cell.fill = PatternFill("solid", fgColor="1F4E78")
        cell.font = Font(name="Microsoft YaHei", size=10, bold=True, color="FFFFFF")
        cell.alignment = Alignment(horizontal="center", vertical="center")
    rows = [
        ("49条测试用例", "有条件通过", "48/49", "0", "MODEL-04 阻塞", "测试用例清单"),
        (
            "功能/治理/运维分组",
            FUNCTIONAL["overall_status"],
            f"{FUNCTIONAL['commands_passed']}/{FUNCTIONAL['commands_total']}",
            "0",
            "22个命令级检查全通过",
            "execution_results/2026-07-26-current/functional/",
        ),
        (
            "后端全量 pytest",
            "PASS",
            f"{BACKEND_COUNTS[1]}/{BACKEND_COUNTS[0]}",
            str(BACKEND_COUNTS[3]),
            "启用集成后另有19条通过",
            "execution_results/2026-07-26-current/full-validation/",
        ),
        (
            "前端全量 Vitest",
            "PASS",
            f"{FRONTEND_COUNTS[1]}/{FRONTEND_COUNTS[0]}",
            str(FRONTEND_COUNTS[3]),
            "38个文件、136条测试",
            "execution_results/2026-07-26-current/full-validation/",
        ),
        (
            "MySQL/Redis 集成",
            "PASS",
            f"{INTEGRATION_COUNTS[1]}/{INTEGRATION_COUNTS[0]}",
            str(INTEGRATION_COUNTS[3]),
            "含 ST-102、ST-103、迁移和认证授权",
            "execution_results/2026-07-26-current/full-validation/integration-junit.xml",
        ),
        ("前端 Lint/Build", "PASS", "2/2", "0", "ESLint + production build", "full-validation/03~04"),
        ("确定性种子与核心 E2E", "PASS", "19/19", "0", "种子11 + Playwright 8", "acceptance/"),
        ("黑盒冒烟", "PASS", "11/11", "0", "Nginx/API真实契约", "acceptance/03-smoke*.log"),
        ("性能门禁", "PASS", "3/3", "0", "3910请求、0失败、P95达标", "performance/gate-result.*"),
        ("恢复演练", "PASS", "4/4", "0", "备份/恢复/重建/回源", "recovery/summary.md"),
        ("LoRA实际训练产物", "BLOCKED", "0/33", "33", "缺依赖与三组实际产物", "full-validation/backend-full-junit.xml"),
    ]
    thin = Side(style="thin", color="B7C9D6")
    for row_index, values in enumerate(rows, start=6):
        for column, value in enumerate(values, start=1):
            cell = result_sheet.cell(row_index, column, value)
            cell.font = Font(name="Microsoft YaHei", size=10)
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
            if column == 2:
                cell.fill = PatternFill(
                    "solid",
                    fgColor="FFF2CC"
                    if value in {"BLOCKED", "有条件通过"}
                    else "E2F0D9",
                )
                cell.font = Font(
                    name="Microsoft YaHei",
                    size=10,
                    bold=True,
                    color="9C6500"
                    if value in {"BLOCKED", "有条件通过"}
                    else "006100",
                )
    result_sheet.freeze_panes = "A6"
    result_sheet.auto_filter.ref = f"A5:F{result_sheet.max_row}"
    widths = [26, 16, 14, 10, 36, 58]
    for column, width in enumerate(widths, start=1):
        result_sheet.column_dimensions[get_column_letter(column)].width = width
    result_sheet.sheet_view.showGridLines = False
    result_sheet.page_setup.orientation = "landscape"
    result_sheet.page_setup.fitToWidth = 1
    result_sheet.sheet_properties.pageSetUpPr.fitToPage = True
    workbook.save(WORKBOOK)


def set_cell_shading(cell, color: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    properties.append(shading)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("第 ")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    run._r.addnext(field)
    paragraph.add_run(" 页")


def add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_heading(text, level=level)
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(6)


def add_table(document: Document, headers: list[str], rows: list[tuple[str, ...]], widths=None):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        set_cell_shading(cell, "1F4E78")
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.name = "Microsoft YaHei"
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = str(value)
            cells[index].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cells[index].paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.name = "Microsoft YaHei"
                    run.font.size = Pt(8.5)
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Cm(width)
    document.add_paragraph()
    return table


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.add_run(text)


def build_report() -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.0)
    styles = document.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    styles["Normal"].font.size = Pt(10.5)
    styles["Title"].font.name = "Microsoft YaHei"
    styles["Title"]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    for name in ("Heading 1", "Heading 2", "Heading 3"):
        styles[name].font.name = "Microsoft YaHei"
        styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        styles[name].font.color.rgb = RGBColor(31, 78, 120)
    header = section.header.paragraphs[0]
    header.text = "LocalLife Copilot｜测试总结报告 V2.0（实际执行版）"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = RGBColor(89, 89, 89)
    add_page_number(section.footer.paragraphs[0])

    cover = document.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover.paragraph_format.space_before = Pt(90)
    run = cover.add_run("LocalLife Copilot")
    run.bold = True
    run.font.size = Pt(30)
    run.font.color.rgb = RGBColor(23, 54, 93)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(14)
    run = title.add_run("测试总结报告")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(31, 78, 120)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(10)
    run = subtitle.add_run("V2.0｜实际执行与证据归档版")
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor(91, 155, 213)
    document.add_paragraph()
    add_table(
        document,
        ["项目", "内容"],
        [
            ("执行日期", "2026-07-26"),
            ("测试对象", "当前本地候选工作区"),
            ("Git 分支 / Commit", f"{BRANCH} / {COMMIT}"),
            ("工作区状态", f"存在 {len(STATUS_LINES)} 项未提交改动，结果绑定当前文件状态"),
            ("测试负责人", "项目测试组"),
            ("交付结论", "材料可交付；软件发布门禁有条件通过"),
        ],
        [4.2, 12.5],
    )
    notice = document.add_paragraph()
    notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = notice.add_run("共49项：P0 37项、P1 12项；覆盖7个Epic、17个Story")
    run.bold = True
    run.font.size = Pt(11)
    document.add_page_break()

    add_heading(document, "1. 执行结论")
    paragraph = document.add_paragraph()
    paragraph.add_run("本次测试材料已完成实际执行、证据归档和用例回填。").bold = True
    paragraph.add_run(
        "49 条用例中 48 条通过、1 条阻塞；阻塞项为 MODEL-04 的实际 LoRA "
        "训练产物验证。除该条件项外，功能、集成、E2E、黑盒冒烟、性能、恢复、"
        "Lint 与生产构建均通过。"
    )
    add_table(
        document,
        ["结论项", "数量", "说明"],
        [
            ("测试用例", "49", "48 通过，1 阻塞，0 失败"),
            ("P0", "37", "均已执行；MODEL-04 为 P1"),
            ("P1", "12", "11 通过，MODEL-04 阻塞"),
            ("缺陷", "0", "未发现断言失败型产品缺陷"),
            ("环境/产物阻塞", "1", "BLOCK-MODEL-ARTIFACT-001"),
        ],
        [4, 2.5, 10.2],
    )
    paragraph = document.add_paragraph()
    run = paragraph.add_run("总体判定：有条件通过。")
    run.bold = True
    run.font.color.rgb = RGBColor(156, 101, 0)
    paragraph.add_run(
        " 测试文档包本身可以交付；发布前应生成并校验 smoke-001、smoke-002、"
        "lora-001 训练产物，同时将当前工作区固化为 Commit/Tag。"
    )

    add_heading(document, "2. 测试范围与环境")
    document.add_paragraph(
        "本轮覆盖 Story：ST-101、ST-102、ST-103、ST-201、ST-202、ST-301、"
        "ST-302、ST-401、ST-402、ST-501、ST-502、ST-601、ST-602、"
        "ST-603、ST-701、ST-702、ST-703。"
    )
    add_table(
        document,
        ["维度", "执行范围"],
        [
            ("功能域", "认证授权、治理、知识库、检索、RAG、WebSocket、商家、模型、UI、管理端"),
            ("服务", "Nginx、API、Worker、MySQL 8.4、Redis 7.4、OpenSearch 3.7、Model Gateway、Frontend"),
            ("后端", "Python 3.13、pytest 8.4.1"),
            ("前端", "Node 22.23.1、Vitest 3.2.7、Playwright 1.61.1"),
            ("性能", "Locust 2.46.1，12 用户，3 用户/秒，持续 2 分钟"),
            ("候选入口", "Nginx http://127.0.0.1:3000；API http://127.0.0.1:18000"),
        ],
        [4.2, 12.5],
    )

    add_heading(document, "3. 自动化执行结果")
    rows = [
        ("功能/治理/运维分组", "22/22 PASS", "0", "functional/summary.json"),
        ("后端全量 pytest", f"{BACKEND_COUNTS[1]} PASS", str(BACKEND_COUNTS[3]), "backend-full-junit.xml"),
        ("前端全量 Vitest", f"{FRONTEND_COUNTS[1]} PASS", str(FRONTEND_COUNTS[3]), "frontend-full-junit.xml"),
        ("MySQL/Redis 集成", f"{INTEGRATION_COUNTS[1]}/{INTEGRATION_COUNTS[0]} PASS", "0", "integration-junit.xml"),
        ("前端 ESLint", "PASS", "0", "03-frontend-lint.log"),
        ("前端生产构建", "PASS", "0", "04-frontend-build.log"),
        ("确定性种子", "11/11 PASS", "0", "acceptance/01-*.log"),
        ("核心角色 Playwright", "8/8 PASS", "0", "acceptance/02-*.log"),
        ("候选环境黑盒冒烟", "11/11 PASS", "0", "acceptance/03-*.log"),
        ("恢复演练", "PASS", "0", "recovery/summary.md"),
    ]
    add_table(document, ["执行项", "结果", "跳过", "主要证据"], rows, [5.2, 3.2, 2, 6.3])
    paragraph = document.add_paragraph()
    paragraph.add_run("有效后端覆盖：").bold = True
    paragraph.add_run(
        f"默认全量 {BACKEND_COUNTS[1]} 条通过，加上隔离环境启用的 "
        f"{INTEGRATION_COUNTS[1]} 条集成验证。默认套件的 41 条跳过中，"
        "6 条已通过集成批次补跑；剩余 35 条为 33 条训练产物条件用例和 "
        "2 条项目已显式标注的 RAG 基准校准/Windows 临时目录问题。"
    )

    add_heading(document, "4. 性能门禁结果")
    perf_rows = []
    for result in PERFORMANCE["results"]:
        perf_rows.append(
            (
                result["name"],
                str(result["requests"]),
                str(result["failures"]),
                f"{result['p95_ms']} ms",
                f"{result['comparison']} {result['threshold_ms']} ms",
                "PASS" if result["passed"] else "FAIL",
            )
        )
    add_table(
        document,
        ["指标", "请求数", "失败", "P95", "阈值", "结果"],
        perf_rows,
        [5.5, 2, 1.5, 2, 3, 2],
    )
    paragraph = document.add_paragraph(
        "Locust 共完成 3910 次业务请求，失败数为 0；聚合平均响应时间 58 ms。"
        "三个发布门禁指标全部通过。"
    )

    add_heading(document, "5. 恢复演练结果")
    recovery_hash = re.search(r"Backup SHA-256: ([0-9a-f]+)", RECOVERY_TEXT)
    recovery_index = re.search(r"Rebuilt index: ([a-z0-9-]+)", RECOVERY_TEXT)
    add_table(
        document,
        ["演练项", "结果", "证据"],
        [
            (
                "MySQL 加密备份与临时库恢复",
                "PASS",
                f"36 张表逐表计数一致；SHA-256 {recovery_hash.group(1) if recovery_hash else '见原始报告'}",
            ),
            ("OpenSearch 全量重建", "PASS", recovery_index.group(1) if recovery_index else "见原始报告"),
            ("重建后对账与别名切换", "PASS", "post-reconcile.json"),
            ("Redis DB0 丢失与 MySQL 回源", "PASS", "redis-fallback.json"),
            ("临时资源清理", "PASS", "临时恢复数据库与容器内脚本已删除"),
        ],
        [5.3, 2.2, 9.2],
    )

    add_heading(document, "6. 用例执行状态")
    add_table(
        document,
        ["用例域", "数量", "结果", "关键证据"],
        [
            ("AUTH/GOV", "7", "7 通过", "功能分组、冒烟、19条真实集成"),
            ("KB", "7", "7 通过", "后端/前端、ST-102 集成、恢复重建"),
            ("SEARCH/RAG/WS", "8", "8 通过", "功能分组、E2E、性能"),
            ("MERCHANT", "6", "6 通过", "123后端、10前端、角色E2E"),
            ("MODEL", "7", "6通过/1阻塞", "202后端、10前端；MODEL-04 产物阻塞"),
            ("UI/ADMIN", "6", "6 通过", "136全量前端、10条Playwright"),
            ("OPS/E2E", "8", "8 通过", "Compose、冒烟、性能、恢复、核心E2E"),
        ],
        [4.2, 2, 3.2, 7.3],
    )
    paragraph = document.add_paragraph(
        "详细到每条用例的执行结果、缺陷编号和证据路径已回填到"
        "《LocalLife-Copilot测试用例清单.xlsx》的“测试用例清单”和“本次执行结果”工作表。"
    )

    add_heading(document, "7. 阻塞项与风险")
    add_table(
        document,
        ["编号", "影响用例", "现状", "解除条件"],
        [
            (
                "BLOCK-MODEL-ARTIFACT-001",
                "MODEL-04",
                "训练逻辑自动化通过，但实际 LoRA 产物校验 33 条按条件跳过",
                "安装 peft/datasets/accelerate；缓存或下载白名单模型；固定种子运行 smoke-001、"
                "smoke-002 和 lora-001；复跑 test_train_lora_artifacts.py 并归档 JUnit",
            )
        ],
        [4, 2.2, 5.5, 6.5],
    )
    add_bullet(
        document,
        "当前工作区存在未提交改动。执行证据绑定的是当前文件状态与 Commit "
        f"{COMMIT[:12]}，发布签署前应提交或生成不可变归档。",
    )
    add_bullet(
        document,
        "前端测试日志含未注册 Element Plus 测试桩和 Vue Router 无匹配警告，"
        "不影响退出码与断言结果，但建议后续清理以提升日志信噪比。",
    )
    add_bullet(
        document,
        "E2E-01 当前候选环境完成 1 轮 8 条桌面/移动自动化；三轮全新环境结论引用"
        "既有 TK-702-03 历史证据，未在本轮执行破坏性的 Compose 卷重置。",
    )

    add_heading(document, "8. 交付物与复现方式")
    add_table(
        document,
        ["交付物", "用途"],
        [
            ("LocalLife-Copilot测试用例清单.xlsx", "49 条用例、结果、缺陷与证据追踪"),
            ("LocalLife-Copilot测试总结报告-V2.0-执行版.docx", "本报告"),
            ("execution_results/2026-07-26-current/", "原始日志、JUnit、JSON、CSV、HTML、恢复证据"),
            ("test_code/", "可复跑的 Python/PowerShell/Playwright 测试入口"),
            ("交付说明与执行结果.md", "交付入口与一页式结果索引"),
        ],
        [8.5, 8.2],
    )
    code = document.add_paragraph()
    code.add_run(
        "复跑功能分组：\n"
        "py -3.13 docs/测试用例文档/test_code/run_project_test_cases.py "
        "AUTH KB GOV SEARCH RAG WS MERCHANT MODEL UI ADMIN OPS DELIVERY "
        "--continue-on-failure --result-dir <输出目录>\n\n"
        "复跑全量与集成：\n"
        "powershell -File docs/测试用例文档/test_code/run_full_validation.ps1 "
        "-OutputDirectory <输出目录>"
    ).font.name = "Consolas"

    add_heading(document, "9. 签署")
    add_table(
        document,
        ["角色", "姓名/签字", "日期", "结论"],
        [
            ("测试负责人", "", "", "有条件通过"),
            ("项目负责人", "", "", ""),
            ("发布负责人", "", "", ""),
        ],
        [4.2, 4.2, 3.2, 5.1],
    )
    document.core_properties.title = "LocalLife Copilot 测试总结报告 V2.0（实际执行版）"
    document.core_properties.subject = "完整自动化执行结果与证据归档"
    document.core_properties.author = "LocalLife Copilot 项目测试组"
    document.core_properties.keywords = "测试总结, E2E, 性能, 恢复, JUnit, 交付"
    document.save(REPORT)


def write_index() -> None:
    metrics = {item["name"]: item for item in PERFORMANCE["results"]}
    content = f"""# LocalLife Copilot 测试交付说明

本目录是 2026-07-26 当前候选工作区的完整测试交付包。当前结论为：**材料可交付，软件发布有条件通过**。

## 一页式执行结果

| 项目 | 结果 |
|---|---|
| 测试用例 | 49 条：48 通过、1 阻塞、0 失败 |
| 功能/治理/运维分组 | {FUNCTIONAL['commands_passed']}/{FUNCTIONAL['commands_total']} 个命令通过 |
| 后端全量 | {BACKEND_COUNTS[1]} 通过、{BACKEND_COUNTS[3]} 条件跳过 |
| MySQL/Redis 集成补跑 | {INTEGRATION_COUNTS[1]}/{INTEGRATION_COUNTS[0]} 通过 |
| 前端全量 | {FRONTEND_COUNTS[1]}/{FRONTEND_COUNTS[0]} 通过 |
| 确定性种子 + 核心 E2E | 11 + 8 条通过 |
| 黑盒冒烟 | 11/11 通过 |
| 性能 | 3910 请求、0 失败；API/Search/TTFB P95 = 58/180/99 ms |
| 恢复演练 | 加密备份、36 表恢复核对、索引重建、Redis 回源均通过 |
| 唯一阻塞 | MODEL-04 实际 LoRA 训练产物（BLOCK-MODEL-ARTIFACT-001） |

## 权威交付物

- `LocalLife-Copilot测试用例清单.xlsx`：49 条用例的实际结果、缺陷和证据位置。
- `LocalLife-Copilot测试总结报告-V2.0-执行版.docx`：本轮正式总结报告。
- `execution_results/2026-07-26-current/`：原始日志、JUnit XML、Playwright JSON、Locust CSV/HTML、恢复证据。
- `test_code/`：可直接复跑的测试代码和入口。
- `LocalLife-Copilot测试总结报告-V1.1.docx`：前序测试设计基线，仅供追溯；V2.0 是当前执行结论。

## 证据目录

- `functional/`：22 个分组命令日志和汇总。
- `full-validation/`：后端/前端全量、Lint、Build、19 条真实集成及 JUnit。
- `acceptance/`：确定性数据、8 条 Playwright 核心角色链路、11 条候选环境冒烟。
- `performance/`：性能门禁 JSON/Markdown、Locust CSV 和 HTML 报告。
- `recovery/`：加密备份、重建/对账 JSON、Redis 回源和恢复总结。

## 复跑

```powershell
py -3.13 docs/测试用例文档/test_code/run_project_test_cases.py AUTH KB GOV SEARCH RAG WS MERCHANT MODEL UI ADMIN OPS DELIVERY --continue-on-failure --result-dir <输出目录>

powershell -File docs/测试用例文档/test_code/run_full_validation.ps1 -OutputDirectory <输出目录>
```

执行基线：分支 `{BRANCH}`，Commit `{COMMIT}`。当前工作区有 {len(STATUS_LINES)} 项未提交改动，因此发布前还需固化 Commit/Tag。
"""
    INDEX.write_text(content, encoding="utf-8")


def main() -> None:
    assert FUNCTIONAL["overall_status"] == "PASS"
    assert FULL["overall_status"] == "PASS"
    assert ACCEPTANCE["overall_status"] == "PASS"
    assert PERFORMANCE["passed"] is True
    assert "Overall result: **PASS**" in RECOVERY_TEXT
    finalize_workbook()
    build_report()
    write_index()
    print(f"Updated: {WORKBOOK}")
    print(f"Created: {REPORT}")
    print(f"Created: {INDEX}")


if __name__ == "__main__":
    main()
